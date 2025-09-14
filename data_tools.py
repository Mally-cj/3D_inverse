#!/usr/bin/env python
# -*- coding:utf-8 -*-
# Power by Zongsheng Yue 2019-09-02 15:11:05
import os
# os.environ['CUDA_VISIBLE_DEVICES'] = '1'
# import cv2
import torch
import pdb
import numpy as np
import random
import matplotlib.pyplot as plt
from matplotlib import gridspec
from skimage.metrics import mean_squared_error, structural_similarity
from skimage.metrics import peak_signal_noise_ratio as PSNR
# from skimage.measure import compare_ssim, compare_psnr, compare_mse
from scipy.signal import butter, filtfilt
from icecream import ic
import sys
import re
from matplotlib.backends.backend_pdf import PdfPages
from contextlib import contextmanager
import time
import matplotlib
from skimage.metrics import peak_signal_noise_ratio as sklearn_psnr
from scipy.stats import pearsonr
# import sklearn
matplotlib.rc("font",family='WenQuanYi Micro Hei')

import threading
from functools import wraps
from typing import List, Optional, Callable, Any
import queue
import time


import multiprocessing as mp
from abc import ABC, abstractmethod


class ThreadRunner:
    def __init__(self):
        self.queue = queue.Queue()
        self.results = []
        self.stop_event = threading.Event()
        self.thread = threading.Thread(target=self.worker)
        self.thread.daemon = True  # 设置为守护线程，主程序结束时自动结束
        self.thread.start()

    def stop(self):
        self.stop_event.set()  # 设置停止标志
        self.queue.put(None)   # 发送停止信号，避免线程阻塞在 queue.get()
        if self.thread.is_alive():
            self.thread.join()

    def run(self, *args, **kwargs):
        self.queue.put((args, kwargs))
        print(f"当前ThreadRunner队列下有 {self.queue.qsize()} 个任务")


    @abstractmethod
    def _run(self, *args, **kwargs): ...

    @abstractmethod
    def _init_worker(self):  # 子类可重写
        pass

    def worker(self):
        self._init_worker()
        while not (self.stop_event.is_set() and self.queue.empty()):
            try:
                item = self.queue.get(timeout=1)  # 添加超时以避免无限阻塞
                if item is None:  # 收到停止信号
                    break
                args, kwargs = item
                self._run(*args, **kwargs)
                self.queue.task_done()
            except queue.Empty:
                continue  # 超时后继续检查停止标志
    
    def wait_end(self):
        ##计时，等待队列中的任务全部执行完毕
        print(f"ThreadRunner等待队列中的任务全部执行完毕")
        start_time = time.time()
        self.queue.join()
        end_time = time.time()
        print(f"ThreadRunner等待队列中的任务全部执行完毕，耗时 {end_time - start_time:.2f} 秒")


class ProcessRunner:
    def __init__(self):
        self.queue = mp.Queue()
        self.results = []
        self.process = mp.Process(target=self.worker)
        self.process.start()

    def stop(self):
        if self.process.is_alive():
            self.process.terminate()
            self.process.join()

    def run(self, *args, **kwargs):
        print("把具体任务提交到队列中，等待执行")
        self.queue.put((args, kwargs))

    @abstractmethod
    def _run(self, *args, **kwargs): ...

    @abstractmethod
    def _init_worker(self):  # 子类可重写
        pass
    
    def worker(self):
        self._init_worker()
        while True:
            args, kwargs = self.queue.get()  # 添加超时以避免无限阻塞
            self._run(*args, **kwargs)



# 全局线程队列实例
# thread_queue = ThreadQueue()


def run_in_queue(task_name: str = None):
    """装饰器：将函数添加到线程队列中排队执行"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            name = task_name or f"{func.__name__}_{int(time.time())}"
            thread_queue.add_task(func, args, kwargs, name)
            return name  # 返回任务名称
        return wrapper
    return decorator


# 兼容旧的装饰器
def run_in_thread(func):
    """装饰器：将函数添加到线程队列中排队执行（兼容旧接口）"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        task_name = f"{func.__name__}_{int(time.time())}"
        thread_queue.add_task(func, args, kwargs, task_name)
        return task_name
    return wrapper


# 为了向后兼容，保留thread_collector接口
class ThreadCollector:
    """兼容旧接口的收集器，实际使用队列系统"""
    
    def join_all(self, timeout: Optional[float] = None):
        """等待所有任务完成"""
        thread_queue.wait_all_done()
    
    def get_alive_count(self) -> int:
        """获取队列中的任务数量"""
        count = thread_queue.get_queue_size()
        if thread_queue.get_current_task():
            count += 1  # 加上正在执行的任务
        return count
    
    def clear(self):
        """清空队列（注意：不会停止正在执行的任务）"""
        while not thread_queue.task_queue.empty():
            try:
                thread_queue.task_queue.get_nowait()
            except queue.Empty:
                break


# thread_collector = ThreadCollector()


class PDFManager:
    def __init__(self, filename):
        self.filename = filename
        self.pdf = PdfPages(self.filename)

    def save_fig_to_pdf(self, fig):
        self.pdf.savefig(fig)
        plt.close(fig)

    def close(self):
        self.pdf.close()


from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io
from docx.shared import Cm  # 导入 Cm 单位

plt.rcParams['font.family'] = 'Times New Roman'
plt.rcParams['font.weight'] = 'bold'
plt.rcParams['mathtext.fontset'] = 'custom'
plt.rcParams['mathtext.rm'] = 'Times New Roman'
matplotlib.rc("font", family='Times New Roman')

class DocxManager:
    def __init__(self, filename):
        self.doc = Document()
        self.filename = filename
        section = self.doc.sections[0]
        section.top_margin = Cm(0.5)  # 顶部边距设为 1厘米
        section.bottom_margin = Cm(0.5)  # 底部边距设为 1厘米
        section.left_margin = Cm(0.5)  # 左侧边距设为 1厘米
        section.right_margin = Cm(0.5)  # 右侧边距设为 1厘米
    def add_heading(self, text,level):
        print(text)
        self.doc.add_heading(text,level)
        
    def add_paragraph(self, text):
        print(text)
        self.doc.add_paragraph(text)

    def add_figure(self, fig, width=8.5):
        # 将Matplotlib图表保存到一个字节流中，而不是文件
        image_stream = io.BytesIO()
        fig.savefig(image_stream, format='png')
        image_stream.seek(0)  # 移动到流的开始位置
        plt.close(fig)
        
        # 将图表添加到Word文档
        self.doc.add_picture(image_stream, width=Inches(width))

    def save(self):
        print(f"save to {self.filename}")
        self.doc.save(self.filename)


@contextmanager
def timermanager(label=''):
    """上下文管理器，用于计时代码块的执行时间"""
    start_time = time.time()
    yield
    end_time = time.time()
    print(f"{label} took {end_time - start_time:.6f} seconds to execute.")


def add_gaussian_band_noise(target_snr, data):
    ##通过输入snr来确定噪声强度
    if target_snr ==0:
        return data,0
    signal_energy = np.linalg.norm(data) ** 2
    noise_energy = signal_energy / (10 ** (target_snr / 10))
    initial_noise = np.random.normal(loc=0, scale=1, size=data.shape)
    noise = filtfilt(np.ones(3)/3, 1, 
                    filtfilt(np.ones(3)/3, 1, initial_noise.T, method='gust').T, method='gust')
    noise = noise * np.sqrt(noise_energy / np.linalg.norm(noise) ** 2)
    noisy_data = data + noise
    actual_snr = 10 * np.log10(signal_energy / np.linalg.norm(noise) ** 2)

    return noisy_data, actual_snr

import numpy as np
import torch
import matplotlib.pyplot as plt

def show_imgs(images, title=None, vmin=0, vmax=0, names=[], showname=False, save_path=None,cmap=plt.cm.jet, docx_manager=None, layershape=None, docx_img_width=8.5):
    '''
    希望输入是batchsize*64*64，这样的，当然也可以batchsize *64*64*1。
    
    '''
    if isinstance(images, dict):
        names = list(images.keys())
        images = list(images.values())
    
    if isinstance(images, list):
        if isinstance(images[0], torch.Tensor):
            images = torch.stack(images)
        else:
            images = np.stack(images)
    elif isinstance(images,torch.Tensor):
        images=images.squeeze().cpu().numpy()
    
    if title is None:
        title = ', '.join(names)
    if vmin == vmax:
        vmax = images.max()
        vmin = images.min()
    
    if layershape is None:
        n_images = min(30, images.shape[0])  # 这里设置为4张图片
        n_rows = int(np.ceil(np.sqrt(n_images)))  # 计算需要的行数
        n_cols = n_rows  # 这里假设是正方形布局
        fig, axs = plt.subplots(n_rows, n_cols, figsize=(8, 8))
    else:
        n_rows, n_cols = layershape
        size = 8
        # fig, axs = plt.subplots(n_rows, n_cols, figsize=(size, size* n_cols),constrained_layout=False)
        fig, axs = plt.subplots(n_rows, n_cols, figsize=(size*n_cols, size),constrained_layout=False)
        
        # print((size, size* n_cols))
    # title = title + f" min={images.min():.4f}, max={images.max():.4f}"
    # fig.suptitle(title) 

    

    for i, ax in enumerate(axs.flat):
        if i < len(images):
            normalized_img = images[i]
            im = ax.imshow(normalized_img, cmap=cmap, vmin=vmin, vmax=vmax)  # 统一使用vmin和vmax
            # ax.axis('off')  # 不显示坐标轴
            if showname:
                ax.set_title(names[i])
            ax.set_xlabel('Trace No.', fontsize=25, fontweight='bold')
            ax.set_ylabel('Time (ms)', fontsize=25, fontweight='bold')
            labels=ax.get_xticklabels() + ax.get_yticklabels()
            for label in labels:
                label.set_fontname('Times New Roman')
                label.set_fontweight('bold')
                label.set_fontsize(19)
        else:
            ax.axis('off')
    cax = fig.add_axes([0.87, 0.1, 0.008, 0.78])  # 添加一个额外的轴用于颜色条

    cbar=plt.colorbar(im, cax=cax,pad=0.001,aspect=10)  # 添加颜色条
    plt.subplots_adjust(left=0.1,right=0.88,wspace=0.001, hspace=0.01)
    cbar.ax.set_title(r'$\mathrm{g/cm^3} \cdot \mathrm{m/s} $', fontsize= 22, fontweight='bold',pad=20)
    cbar.ax.tick_params(labelsize=21,pad=10)
    if save_path is not None:
        np.save(os.path.join(save_path, title + '.npy'), images)
        plt.savefig(fname=os.path.join(save_path, title + '.jpg'), format='jpg', bbox_inches='tight', dpi=600)
    else:
        fig.suptitle(title, fontsize=40)

    if docx_manager is not None:
        fig = plt.gcf() 
        docx_manager.add_figure(fig, width=docx_img_width)
    plt.show()

# show_imgs([ddpm_img-true,supervise-true,unsupervise-true,pd_img-true],layershape=(1,4),vmin=-800,vmax=800,save_path=save_path,title='nonlinear_0_0.012_30_cha_compare4')


import matplotlib.pyplot as plt
import os

def show_plots(values, x_tags=[3, 5, 6, 9, 10, 12, 15], save=False, xlabel='PSNR(dB)', ylim=None, save_dir=None, title='psnr_compare', x_colors=None, markers=None, marker_sizes=None):
    """
    绘制图表并支持自定义横坐标标签颜色、标记符号及其大小。

    参数:
    - values: 数据字典，键为算法名称，值为对应的纵坐标值列表。
    - x_tags: 横坐标标签列表，默认为 [3, 5, 6, 9, 10, 12, 15]。
    - save: 是否保存图表，默认为 False。
    - xlabel: 横坐标标签名称，默认为 'PSNR(dB)'。
    - ylim: 纵坐标范围，默认为 None。
    - save_dir: 图表保存路径，默认为 None。
    - title: 图表标题，默认为 'psnr_compare'。
    - x_colors: 横坐标颜色字典，键为横坐标值，值为颜色，默认为 None。
    - markers: 标记符号列表，默认为 None（使用默认标记符号）。
    - marker_sizes: 标记符号大小列表，默认为 None（使用默认大小 8）。
    """
    # 设置字体
    plt.rcParams['font.sans-serif'] = ['Times New Roman', 'SimHei']  # 支持中文
    plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

    plt.figure(figsize=(10, 5))  # 调整图表大小

    # 默认标记符号
    default_markers = ['o', 's', '^', 'D', 'v', 'p', '*', 'X', '<', '>']
    if markers is None:
        markers = default_markers  # 如果未传入 markers，则使用默认值

    # 默认标记符号大小
    if marker_sizes is None:
        marker_sizes = [8] * len(values)  # 如果未传入 marker_sizes，则使用默认大小 8

    marker_index = 0  # 标记符号索引

    for (algorithm, psnrs), marker_size in zip(values.items(), marker_sizes):
        # 为每条线设置不同的标记符号和大小
        plt.plot(x_tags, psnrs, marker=markers[marker_index % len(markers)], markersize=marker_size, label=algorithm)
        marker_index += 1  # 切换到下一个标记符号

    plt.legend(loc='lower right', bbox_to_anchor=(1.0, 0), prop={'weight': 'bold', 'size': 14})  # 调整图例位置

    plt.xlabel('Frequency (Hz)', fontsize=17, fontweight='bold')
    plt.ylabel(xlabel, fontsize=17, fontweight='bold')

    # 设置横坐标标签
    plt.xticks(x_tags, fontsize=15, fontweight='bold')
    
    # 设置横坐标标签颜色
    if x_colors is not None:
        for i, label in enumerate(plt.gca().get_xticklabels()):
            if x_tags[i] in x_colors:  # 如果横坐标值在 x_colors 字典中
                label.set_color(x_colors[x_tags[i]])  # 设置颜色
            else:
                label.set_color('black')  # 默认颜色为黑色
    
        # 动态设置纵坐标范围
    if ylim is None:
        all_values = [value for psnrs in values.values() for value in psnrs]  # 展平所有纵坐标值
        min_val, max_val = min(all_values), max(all_values)
        margin = (max_val - min_val) * 0.1  # 增加 10% 的间距
        plt.ylim(min_val - margin, max_val + margin)
    else:
        plt.ylim(ylim[0], ylim[1]) 
    plt.grid(True)
    
    if save and save_dir is not None:
        # 检查保存路径是否存在
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)
        save_path = os.path.join(save_dir, title + '.jpg')
        plt.savefig(fname=save_path, format='jpg', bbox_inches='tight', dpi=600)
    
    plt.show()

# 示例调用
values = {
    'Algorithm1': [30, 32, 34, 36, 38, 40, 42],
    'Algorithm2': [28, 30, 32, 34, 36, 38, 40],
    'Algorithm3': [29, 31, 33, 35, 37, 39, 41]
}

# # 定义横坐标颜色
# x_colors = {
#     5: 'blue',  # 横坐标为 5 的标签为蓝色
#     10: 'red'   # 横坐标为 10 的标签为红色
# }

# # 定义自定义标记符号
# custom_markers = ['o', 's', '^']  # 圆形、正方形、三角形

# # 定义自定义标记符号大小
# custom_marker_sizes = [10, 12, 14]  # 每条线的标记符号大小

# # 调用函数
# show_plots(values, x_colors=x_colors, markers=custom_markers, marker_sizes=custom_marker_sizes)



def show_plots1(values,x_tags=[3, 5,6, 9, 10,12, 15],save=False,xlabel='PSNR(dB)',ylim=None,save_dir=None,title='psnr_compare'):

    plt.rcParams['font.sans-serif'] = ['Times New Roman']
    # 低频情况
    # x_tags = [3, 5,6, 9, 10,12, 15]

    plt.figure(figsize=(10, 5))  # 可以调整图表大小
    for algorithm, psnrs in values.items():
        plt.plot(x_tags, psnrs, marker='o', label=algorithm)

    plt.legend(loc='lower right',bbox_to_anchor=(-2.0,0),prop={'weight': 'bold', 'size': 14})

    # plt.xlabel('Low Frequency (Hz)', fontsize=17,fontweight='bold')
    plt.xlabel('Frequency (Hz)', fontsize=17,fontweight='bold')

    plt.ylabel(xlabel, fontsize=17,fontweight='bold')
    plt.xticks(x_tags,fontsize=15,fontweight='bold')
    plt.yticks(fontsize=15,fontweight='bold')
    
    if ylim is not None:
        plt.ylim(ylim[0], ylim[1]) 
    plt.grid(True)
    # save=True
    if save and save_dir is not None:
        # save_dir='/home/shendi_gjh_cj/codes/double_dps/pic/tem/'
        save_path=save_dir+title+'.jpg'
        check_file_exist(save_path)
        # np.save(save_dir+title+'.npy',img)
        # plt.savefig(fname=save_path,dpi=600, bbox_inches='tight', format='jpeg')
        plt.savefig(fname=save_path,format='jpg', bbox_inches='tight', dpi=600)
    plt.show()




def show2(images):
    assert(images.shape[1]==images.shape[2])
    # images = np.transpose(images, (2, 0, 1))
        # 定义要显示的图片数量和布局
    n_images = min(25,images.shape[0])  # 这里设置为4张图片
    n_rows = int(np.ceil(np.sqrt(n_images)))  # 计算需要的行数
    n_cols = n_rows  # 这里假设是正方形布局
    # 假设 images 是一个numpy数组，其中包含多张图像数据
    # 其他代码保持不变...

    fig, axs = plt.subplots(n_rows, n_cols, figsize=(5, 5), sharex=True, sharey=True)
    fig.suptitle("jddd")

    global_min = np.min(images)
    global_max = np.max(images)

    for i, ax in enumerate(axs.flat):
        if i < len(images):
            # 显示图片并归一化
            normalized_img = (images[i] - global_min) / (global_max - global_min)
            im = ax.imshow(normalized_img, cmap=plt.cm.seismic, vmin=global_min, vmax=global_max)  # 确保所有子图使用相同的颜色范围
            ax.axis('off')  # 不显示坐标轴
        else:
            ax.axis('off')  # 隐藏多余的子图

    # 添加颜色轴到图像右侧
    cbar = fig.colorbar(im, ax=axs.ravel().tolist(), orientation='vertical', shrink=0.6)  # axs.ravel().tolist()使颜色轴与所有子图共享
    cbar.ax.set_ylabel('Colorbar Label')  # 设置颜色轴的标签

    plt.tight_layout()
    plt.show()



def cat_array(arr,detail="caculate:"):
    print(detail)
    if isinstance(arr,torch.Tensor):
        print(f"arr.mean={arr.mean()}")
        print(f"arr.var={arr.var()}")
        print(f"arr.min={arr.min()}")
        print(f"arr.max={arr.max()}")
    else:
        print(f"arr.mean={np.mean(arr)}")
        print(f"arr.var={np.var(arr)}")
        print(f"arr.min={np.min(arr)}")
        print(f"arr.max={np.max(arr)}")


def read_sgy_road(filename,idx_road=99):
    import obspy
    stream=obspy.read(filename)
    tr = stream[idx_road]  # 获取第一个 Trace
    return tr.data

def count_ssim_mse(true_impedance,inverted_impedance,title='img',show=True):
    inverted_impedance=inverted_impedance.reshape(true_impedance.shape)
    if show:
        print(f"==============={title}====================")
        print(f"Max:{inverted_impedance.max()}")
        print(f"Min:{inverted_impedance.min()}")
        print(f"Mean:{inverted_impedance.mean()}")
        print(f"Mean:{inverted_impedance.var()}")
        mse = mean_squared_error(true_impedance, inverted_impedance)
        print(f"MSE: {mse}")
        ssim = structural_similarity(true_impedance, inverted_impedance,data_range=1)
        print(f"SSIM: {ssim}")
        print("===================================")



def plot_ssim_curves(dB_values, ssim_data_dict):
    """
    绘制不同噪声水平下的多条SSIM曲线。
    
    参数:
    dB_values: 一个列表，包含不同的噪声水平（dB）。
    ssim_data_dict: 一个字典，键为曲线的标签，值为对应每个噪声水平的SSIM值列表。
    """
    plt.figure(figsize=(10, 6))

    # 遍历字典，为每组数据绘制一条曲线
    for label, ssim_data in ssim_data_dict.items():
        # 确保每个数据集长度与噪声水平列表相同
        if len(ssim_data) != len(dB_values):
            raise ValueError(f"数据长度不匹配：{label} 的数据长度与噪声水平数量不符。")
        plt.plot(dB_values, ssim_data, label=label, marker='o')  # 这里使用默认的标记，可以根据需要调整
    
    plt.title('SSIM Values at Different Noise Levels (dB)')
    plt.xlabel('Noise Level (dB)')
    plt.ylabel('SSIM')
    
    plt.legend()
    plt.grid(True)
    plt.show()


# def single_imshow(img, title="img", vmin=0, vmax=0, cmap=plt.cm.jet, save=False,
#                  dpi=600, save_dir='tmp', pdf_manager=None, docx_manager=None,
#                  docx_img_width=4.0, arrow_details=None, alpha_loc=None, 
#                  y_tickets=None, x_tickets=None, y_range=None):
#     plt.rcParams['font.family'] = 'Times New Roman'
#     plt.rcParams['font.weight'] = 'bold'
#     plt.rcParams['mathtext.fontset'] = 'custom'
#     plt.rcParams['mathtext.rm'] = 'Times New Roman'
#     matplotlib.rc("font", family='Times New Roman')
    
#     if vmin == vmax:
#         vmin = img.min()
#         vmax = img.max()
#     print(f"vmin={vmin}")
#     print(f"vmax={vmax}")
#     assert img.max() >= vmin
#     assert img.min() <= vmax
    
#     if isinstance(img, torch.Tensor):
#         img = img.squeeze().cpu().numpy()
    
#     fig, ax = plt.subplots()
    
#     # 关键修改：根据 y_range 设置 extent
#     if y_range is not None:
#         ymin, ymax = y_range
#         extent = [0, img.shape[1], ymax, ymin]  # x从0到列数，y从ymax到ymin（适配imshow原点在左上角）
#     else:
#         extent = None
    
#     cax = ax.imshow(img, cmap=cmap, vmin=vmin, vmax=vmax, aspect='auto', extent=extent)
#     cbar = fig.colorbar(cax, ax=ax, pad=0.01, aspect=60, shrink=0.85)
#     cbar.ax.set_title(r'$\mathrm{g/cm^3} \cdot \mathrm{m/s} $', fontsize=12, fontweight='bold')
    
#     for label in cbar.ax.get_yticklabels():
#         label.set_fontname('Times New Roman')
#         label.set_fontweight('bold')
#         label.set_fontsize(12)
    
#     ax.set_xlabel('Trace No.', fontsize=17, fontweight='bold')
#     ax.set_ylabel('Time (ms)', fontsize=17, fontweight='bold')
    
#     # 设置纵坐标刻度和标签
#     if y_range is not None:
#         ymin, ymax = y_range
#         if y_tickets is None:
#             y_tickets = np.linspace(ymin, ymax, 6)  # 默认生成6个均匀分布的刻度
#         ax.set_yticks(np.linspace(ymin, ymax, len(y_tickets)))  # 根据y_range映射刻度位置
#         # ax.set_yticklabels([f"{int(tick)}" for tick in y_tickets], fontweight='bold')    ##坐标轴显示整数
#         ax.set_yticklabels([f"{tick:.2f}" for tick in y_tickets], fontweight='bold')    #坐标轴显示小数

#         # ax.set_ylim(ymin, ymax)  # 强制设置纵坐标范围
#     else:
#         if y_tickets is None:
#             # y_tickets = np.arange(0, img.shape[0], img.shape[0] // 6)
#             y_tickets = np.arange(0, img.shape[0],0.05)
#         ax.set_yticks(y_tickets)
#         ax.set_yticklabels(y_tickets, fontweight='bold')
    
#     # 设置横坐标
#     if x_tickets is None:
#         x_tickets = np.arange(0, img.shape[1], 20)
#     ax.set_xticks(x_tickets)
#     ax.set_xticklabels(x_tickets, fontweight='bold')
    
#     ax.tick_params(axis='both', which='major', labelsize=16)
    
#     if arrow_details is not None:
#         for point in arrow_details:
#             ax.annotate(
#                 '',
#                 xy=point['end'],
#                 xytext=point['start'],
#                 arrowprops=dict(facecolor='black', edgecolor='black', shrink=0.05, width=2, headwidth=8, headlength=10),
#             )
    
#     if alpha_loc is not None:
#         cax.set_alpha(alpha_loc)
    
#     if save:
#         save_path = os.path.join(save_dir, title + '.jpg')
#         check_file_exist(save_path)
#         np.save(os.path.join(save_dir, title + '.npy'), img)
#         plt.savefig(fname=save_path, format='jpg', bbox_inches='tight', dpi=dpi)
#     else:
#         title = title + f" var={np.var(img):.4f},mean={np.mean(img):.4f},min={np.min(img):.4f},max={np.max(img):.4f}"
#         plt.title(title)
    
#     if docx_manager is not None:
#         fig = plt.gcf() 
#         docx_manager.add_figure(fig, width=docx_img_width)
#     if pdf_manager is not None:
#         pdf_manager.save_fig_to_pdf(fig)
#     plt.show()
# y_tickets=None
# y_range = (0.8, 0.95)  # 纵坐标显示范围
# single_imshow(record_dealed['true'][:-32],vmin=-1,vmax=1,cmap=plt.cm.seismic,title='true_record',
            #   save=False,alpha_loc=alpha_loc,save_dir=save_dir,y_tickets=y_tickets,y_range=y_range)
def single_imshow(img,title="img",vmin=None,vmax=None,cmap=plt.cm.jet,save=False,
                  dpi=600,save_dir='tmp',pdf_manager=None,docx_manager=None,
                  docx_img_width=4.0,arrow_details=None,alpha_loc=None,y_tickets=None,
                  x_tickets=None):
    if vmin==None:
        vmin=img.min()
    if vmax==None:
        vmax=img.max()
    plt.rcParams['font.family'] = 'Times New Roman'
    plt.rcParams['font.weight'] = 'bold'
    plt.rcParams['mathtext.fontset'] = 'custom'
    plt.rcParams['mathtext.rm'] = 'Times New Roman'
    matplotlib.rc("font", family='Times New Roman')
    if vmin==vmax:
        vmin=img.min()
        vmax=img.max()
    print(f"vmin={vmin}")
    print(f"vmax={vmax}")
    assert(img.max()>=vmin)
    assert(img.min()<=vmax)

    if isinstance(img,torch.Tensor):
        img=img.squeeze().cpu().numpy()
    fig, ax = plt.subplots()
    cax = ax.imshow(img, cmap=cmap,vmin=vmin,vmax=vmax,aspect='auto',)
    ##shrink控制颜色条的长度相对于图像高度的比例
    # cbar = fig.colorbar(cax, ax=ax,pad=0.01,aspect=50,shrink=0.85)
    # 根据图片长宽比自适应调整颜色条
    img_height, img_width = img.shape
    aspect_ratio = img_width / img_height
    
    # 根据长宽比动态调整颜色条参数
    if aspect_ratio > 2:  # 宽图
        fraction = 0.03
        shrink = 0.9
    elif aspect_ratio < 0.5:  # 高图
        fraction = 0.06
        shrink = 0.6
    else:  # 接近正方形的图
        fraction = 0.046
        shrink = 0.8
    cbar = fig.colorbar(cax, ax=ax, aspect=30, shrink=1.0, pad=0.01)
    # cbar = fig.colorbar(cax, ax=ax, pad=0.01, fraction=fraction, shrink=shrink)

    # from mpl_toolkits.axes_grid1 import make_axes_locatable
    # divider = make_axes_locatable(ax)
    # cax = divider.append_axes("right", size="5%", pad=0.1)  # 调整 size 和 pad 控制宽度和间距
    # cbar = fig.colorbar(cax, cax=cax)


    # cbar.ax.set_title(r'$\mathrm{g/cm^3} \cdot \mathrm{m/s} $', fontsize= 12, fontweight='bold')
    for label in cbar.ax.get_yticklabels():
        label.set_fontname('Times New Roman')
        label.set_fontweight('bold')
        label.set_fontsize(12)
        
    ax.set_xlabel('Trace No.', fontsize=17, fontweight='bold')
    ax.set_ylabel('Time (ms)', fontsize=17, fontweight='bold')
    
    if y_tickets is None:
        number=img.shape[0]//6
        if number>10:
            number-=number%10
        y_tickets=np.arange(0, img.shape[0], number)
    if x_tickets is None:
        number2=img.shape[1]//6
        if number2>10:
            number2-=number2%10
        x_tickets=np.arange(0, img.shape[1], number2)
    
    ax.set_xticks(x_tickets)  # 横坐标每隔10个格显示一次
    ax.set_yticks(y_tickets)  # 纵坐标每隔10个格显示一次

    ax.set_xticklabels(x_tickets,fontweight='bold')
    ax.set_yticklabels(y_tickets,fontweight='bold')
    
    ax.tick_params(axis='both', which='major', labelsize=16)  # 设置刻度标签的大小和加粗
    if arrow_details is not None:
        for point in arrow_details:
               ax.annotate(
                '',  # 空文本
                xy=point['end'],  # 箭头指向的点 (x, y)
                xytext=point['start'],  # 箭头起始点 (x, y)
                arrowprops=dict(facecolor='black', edgecolor='black', shrink=0.05, width=2, headwidth=8, headlength=10),  # 黑色箭头
            )
    if alpha_loc is not None:
        cax.set_alpha(alpha_loc)

    if save:
        # save_dir='/home/shendi_gjh_cj/codes/double_dps/pic/tem/'
        save_path=os.path.join(save_dir,title+'.jpg')
        check_file_exist(save_path)
        # np.save(os.path.join(save_dir,title+'.npy'),img)
        plt.title(title)

        plt.savefig(fname=save_path,format='jpg', bbox_inches='tight', dpi=dpi)
    else:
        title=title+f" var={np.var(img):.4f},mean={np.mean(img):.4f},min={np.min(img):.4f},max={np.max(img):.4f}"
        plt.title(title)
    if docx_manager is not None:
        fig = plt.gcf() 
        # docx_manager.add_figure(fig,width=docx_img_width)
        docx_manager.add_figure(fig)

    if pdf_manager is not None:
        pdf_manager.save_fig_to_pdf(fig)
    title=title+f" var={np.var(img):.4f},mean={np.mean(img):.4f},min={np.min(img):.4f},max={np.max(img):.4f}"
    # plt.title(title)
    plt.show()





# def single_imshow(img, title="img", vmin=0, vmax=0, cmap=plt.cm.jet, save=False,
#                   dpi=600, save_dir='tmp', pdf_manager=None, docx_manager=None,
#                   docx_img_width=4.0, arrow_details=None, alpha_loc=None, y_tickets=None,
#                   x_tickets=None):
#     plt.rcParams['font.family'] = 'Times New Roman'
#     plt.rcParams['font.weight'] = 'bold'
#     plt.rcParams['mathtext.fontset'] = 'custom'
#     plt.rcParams['mathtext.rm'] = 'Times New Roman'
#     matplotlib.rc("font", family='Times New Roman')
    
#     if vmin == vmax:
#         vmin = img.min()
#         vmax = img.max()
    
#     print(f"vmin={vmin}")
#     print(f"vmax={vmax}")
#     assert(img.max() >= vmin)
#     assert(img.min() <= vmax)

#     if isinstance(img, torch.Tensor):
#         img = img.squeeze().cpu().numpy()
    
#     fig, ax = plt.subplots()
#     cax = ax.imshow(img, cmap=cmap, vmin=vmin, vmax=vmax)
    
#     cbar = fig.colorbar(cax, ax=ax, pad=0.01, aspect=50, shrink=0.85)
#     cbar.ax.set_title(r'$\mathrm{g/cm^3} \cdot \mathrm{m/s} $', fontsize=12, fontweight='bold')
    
#     for label in cbar.ax.get_yticklabels():
#         label.set_fontname('Times New Roman')
#         label.set_fontweight('bold')
#         label.set_fontsize(12)
        
#     ax.set_xlabel('Trace No.', fontsize=17, fontweight='bold')
#     ax.set_ylabel('Time (ms)', fontsize=17, fontweight='bold')
    
#     # If y_tickets is None, create a default range for the y-ticks
#     if y_tickets is None:
#         number = img.shape[0] // 3
#         if number > 10:
#             number -= number % 10
#         y_tickets = np.arange(0, img.shape[0], number)
#         ic(len(y_tickets))
#         ic(number)
#     # Apply a linear transformation to the y-tick labels
#     # transformed_labels = y_tickets * (vmax - vmin) / img.shape[0]
#     transformed_labels = y_tickets *0.002+0.78


#     # If x_tickets is None, create a default range for the x-ticks
#     if x_tickets is None:
#         number2 = img.shape[1] // 6
#         if number2 > 10:
#             number2 -= number2 % 10
#         x_tickets = np.arange(0, img.shape[1], number2)
    
#     ax.set_xticks(x_tickets)
#     ax.set_yticks(y_tickets)
    
#     ax.set_xticklabels(x_tickets, fontweight='bold')
#     ax.set_yticklabels([f'{label:.2f}' for label in transformed_labels], fontweight='bold')
    
#     ax.tick_params(axis='both', which='major', labelsize=16)

#     if arrow_details is not None:
#         for point in arrow_details:
#             ax.annotate(
#                 '',  # Empty text
#                 xy=point['end'],  # Arrow endpoint (x, y)
#                 xytext=point['start'],  # Arrow start point (x, y)
#                 arrowprops=dict(facecolor='black', edgecolor='black', shrink=0.05, width=2, headwidth=8, headlength=10),
#             )
    
#     if alpha_loc is not None:
#         cax.set_alpha(alpha_loc)

#     if save:
#         save_path = os.path.join(save_dir, title + '.jpg')
#         # check_file_exist(save_path)  # Assuming you have this function
#         np.save(os.path.join(save_dir, title + '.npy'), img)
#         plt.savefig(fname=save_path, format='jpg', bbox_inches='tight', dpi=dpi)
#     else:
#         title = title + f" var={np.var(img):.4f},mean={np.mean(img):.4f},min={np.min(img):.4f},max={np.max(img):.4f}"
#         plt.title(title)
    
#     if docx_manager is not None:
#         fig = plt.gcf()
#         docx_manager.add_figure(fig, width=docx_img_width)
    
#     if pdf_manager is not None:
#         pdf_manager.save_fig_to_pdf(fig)
    
#     title = title + f" var={np.var(img):.4f},mean={np.mean(img):.4f},min={np.min(img):.4f},max={np.max(img):.4f}"
#     plt.show()
# # y_tickets=np.arange(0.72, 0.99, 0.05)
# y_tickets=None
# single_imshow(record_dealed['true'][:-32],vmin=-1,vmax=1,
#               cmap=plt.cm.seismic,title='true',save=False,
#               y_tickets=y_tickets
#               )


def phaseshift(w, d): #相位旋转, w可以用估计一个常相位子波
    from scipy.fftpack import fft, ifft
    sizew=len(w)
    wf_shift = fft(w)*np.exp((1j)*(3.1415926*d/180))    
    wf_shift[int((sizew+1)/2)-1:sizew] = np.conj(wf_shift[int((sizew+1)/2):0:-1])
    w_shift = np.real(ifft(wf_shift))    
    return w_shift




def cat_hist(x,threshold=None,min=0,max=1,title="Image Histogram",bar_cnt=50,win=0):
    '''
    threshold:小于这个值的像素点的数量
    limit1:限制小数点的位数,这是因为存在两个值小数值很临近。
    list_cnt:取前多少个值
    winw:窗口的大小
    '''

    wait_sort=x.flatten()
    if threshold is not None:
        num_small_values = torch.sum(wait_sort< threshold).item()
        num_small_values_rate=num_small_values/wait_sort.numel()
        print(f"{num_small_values} values are smaller than {threshold},\
              the rate is {num_small_values_rate}")
    if win >0:
        min=wait_sort.mean()-win/2
        max=min+win
    
    unique_pixels = np.unique(wait_sort, axis=0)
    sort_out=np.sort(unique_pixels)
    sort_out=np.clip(sort_out,min,max)

    print(f"unique_pixel 's number is {len(unique_pixels)}")
    print(f"total_pixel 's number is {len(wait_sort)}")

    mean=wait_sort.mean()
    var=wait_sort.var()

    bin_edges = np.linspace(sort_out[0], sort_out[-1], num=bar_cnt)  # 创建20个桶
    bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2  # 计算每个桶的中心点
    bin_width = bin_edges[1] - bin_edges[0]  # 计算桶宽度

    counts,_= np.histogram(wait_sort, bins=bin_edges)
    frequency = counts /counts.sum()  # 计算频率
    cumulative_frequency = np.cumsum(frequency)

    plt.bar(bin_centers, frequency,  edgecolor='black',width=bin_width)  # 绘制直方图
    plt.xlabel('Data values')
    plt.ylabel('Frequency')
    plt.title(f'{title} mean={mean:.4f},var={var:.4f},min={unique_pixels.max():.5f},max={unique_pixels.max():.5f}')
    plt.show()
    

    # plt.bar(bin_centers, cumulative_frequency,  edgecolor='black',width=bin_width)  # 绘制直方图
    # plt.xlabel('Data values')
    # plt.ylabel('cumulative_frequency')
    # plt.title(f'{title} mean={mean:.4f},var={var:.4f},min={unique_pixels.max():.5f},max={unique_pixels.max():.5f}')
    # plt.show()

# x=np.random.randn(64,64)
# cat_hist(x,win=5)


import matplotlib.pyplot as plt
import torch

def compare_single_road(imgs, labels=[], idx=0, title=None, docx_manager=None, docx_img_width=4.0, axis=0, y_min=None, y_max=None):
    # assert(len(imgs)==len(labels))
    if title is None:
        title = f"compare road:{idx} from {len(imgs)} imgs"
    if isinstance(imgs, dict):
        labels = list(imgs.keys())
        imgs = list(imgs.values())
    else:  ##否则认为img和labels
        if len(labels) < len(imgs):
            cha = len(imgs) - len(labels)
            labels = ['img' + str(i) for i in range(cha)]

    plt.figure()
    for i in range(len(imgs)):
        if isinstance(imgs[i], torch.Tensor):
            if axis == 0:
                plt.plot(imgs[i].cpu()[:, idx], label=labels[i])
            else:
                plt.plot(imgs[i].cpu()[idx, :], label=labels[i])
        else:
            if axis == 0:
                plt.plot(imgs[i][:, idx], label=labels[i])
            else:
                plt.plot(imgs[i][idx, :], label=labels[i])

    plt.title(title)
    plt.legend()

    # 设置纵坐标范围
    if y_min is not None and y_max is not None:
        plt.ylim(y_min, y_max)

    if docx_manager is not None:
        fig = plt.gcf()
        docx_manager.add_figure(fig, width=docx_img_width)  
# imgs=[]
# labels=[]
# for a in range(3):
#     imgs.append(np.random.rand(128,128))
#     labels.append(str(a))
# compare_single_road(imgs,labels)

def check_tensor_tocpu(x,to_numpy=False):
    if isinstance(x, torch.Tensor):    
        x=x.cpu()
        if to_numpy:
            x=x.numpy() 
    return x

def calculate_dif(img1,img2,show=False):
    """
    计算每道的距离差的总和
    """
    img1=check_tensor_tocpu(img1,to_numpy=True)
    img2=check_tensor_tocpu(img2,to_numpy=True)
    difs=np.mean(img1-img2,axis=0)
    mean=difs.mean()
    var=difs.var()
    print(f"mean={mean},var={var}")
    print(f"min={difs.min()},max={difs.max()}")
    
    return mean,var,difs
    # print(mean)    
    # print(difs.shape)
    
    

def dipin(speed,dipin0=0.012):
    # dipin0=0.012
    BB,AA= butter(2,dipin0,'low')
    nsmoothz,nsmoothx=20,20
    mback= filtfilt(BB,AA,speed.T).T
    mback = filtfilt(np.ones(nsmoothz)/float(nsmoothz),1, mback,axis=0)
    mback = filtfilt(np.ones(nsmoothx)/float(nsmoothx),1, mback, axis=1)
    return mback,dipin0


def calculate_psnr_ssim(trueimg, img,datarange=4700-1027):
    # import cv2
    
    """
    psnr -- 两个图像之间的峰值信噪比。
    ssim -- 两个图像之间的结构相似性指数。
    """
    # trueimg=trueimg.astype(img.dtype)
    # # img1=img1.sequeeze()
    # trueimg=np.squeeze(trueimg,axis=None)
    # img=np.squeeze(img,axis=None)

    ssim = structural_similarity(trueimg,img,data_range=datarange)
    psnr=sklearn_psnr(img, trueimg.squeeze(), data_range=trueimg.max())
    pearsonr_v=pearsonr(img.flatten(), trueimg.flatten()).statistic
    mp={
        'psnr':psnr,
        'ssim':ssim,
        'peasonr':pearsonr_v,
    }
    return mp
##test
# import matplotlib.pyplot as plt
# plt.rcParams['font.sans-serif'] = ['Times New Roman']  # 用来正常显示中文标签
# plt.rcParams['axes.unicode_minus'] = False # 用来正常显示负号
# plt.rcParams['xtick.direction'] = 'in'#将x周的刻度线方向设置向内
# plt.rcParams['ytick.direction'] = 'in'#将y轴的刻度方向设置向内
# single_imshow(np.random.rand(64, 64))
# import matplotlib    
# print(matplotlib.matplotlib_fname())


def check_file_exist(file_path,sure=False):
    import os
    if os.path.exists(file_path):
        # 提示用户并获取用户输入
        if sure is False:
            user_choice = input(f"文件 {file_path} 已存在。是否删除它？(yes/no): ").lower()
        else:
            user_choice="yes"
            
        if user_choice == "yes":
            # 尝试删除文件
            try:
                os.remove(file_path)
                print("文件已成功删除，程序将继续运行。")
            except Exception as e:
                print(f"删除文件时发生错误：{e}")
                exit(1)  # 程序异常终止
        elif user_choice == "no":
            print("用户选择不删除文件，程序终止。")
            exit(0)  # 正常终止程序
        else:
            print("无效的输入，请输入 'yes' 或 'no'。")
            exit(1)  # 输入无效，终止程序
    else:
        print("文件不存在，程序将创建新文件或进行后续操作。")
        


def tensorboard_save_model(model,intensor,logdir=".",):
    from tensorboardX import SummaryWriter
    with SummaryWriter(log_dir=logdir) as sw:  # 实例化 SummaryWriter ,可以自定义数据输出路径
        sw.add_graph(model,intensor)  # 输出网络结构图
        sw.close()  # 关闭  sw
        


def showtensor_RGB(x,title="rgb"):
    '''
    x: 3*w*c的tensor
    '''
    x=x.squeeze()
    if x.shape[0]<=3:
        x=x.permute(1,2,0)
    x=x.cpu().numpy()
    x = ((x - x.min()) * (1/(x.max() - x.min()) * 255)).astype('uint8')
    plt.figure()
    plt.title(title)
    plt.imshow(x)


def rgb_to_gray(img):
    img=img.squeeze()
    c_gray=img[0]*0.299+img[1]*0.587+img[2]*0.114
    return c_gray



def instantiate_from_config(config): 
    import importlib
    def get_obj_from_str(string, reload=False):
        module, cls = string.rsplit(".", 1)
        if reload:
            module_imp = importlib.import_module(module)
            importlib.reload(module_imp)
        return getattr(importlib.import_module(module, package=None), cls)

    if not "target" in config:
        if config == '__is_first_stage__':
            return None
        elif config == "__is_unconditional__":
            return None
        raise KeyError("Expected key `target` to instantiate.")
    return get_obj_from_str(config["target"])(**config.get("params", dict()))



import os, sys, glob
def load_model(folder,epoch1=0,ckpt=None,argsconfig=None):
    dir_folder="/home/shendi_gjh_cj/codes/latent-diffusion/"

    if ckpt is None:
        if epoch1==0:
            ckpt=folder+f"/checkpoints/epoch=*.ckpt"
        else:
            ckpt=folder+f"/checkpoints/epoch={epoch1:06d}*.ckpt"
        ckpt=os.path.join(dir_folder,ckpt)


    
    from omegaconf import OmegaConf
    # from main import instantiate_from_config
    # config_file=folder+"/configs/*project.yaml"

    config_file=os.path.join(dir_folder,folder+"/configs/*project.yaml")
    print(f"config_file is{config_file}")
    config_file=glob.glob(config_file)[0]

    ckpt = glob.glob(ckpt)[0]
    print(f"load ckpt from {ckpt}")


    device = torch.device("cuda") if torch.cuda.is_available() else torch.device("cpu")
    print(device)   
    config = OmegaConf.load(config_file)
    model = instantiate_from_config(config.model).to(device)
    model.load_state_dict(torch.load(ckpt,map_location=device)["state_dict"],
                            strict=False)
    print(device)

    mp={}
    if ckpt is None:
        epoch = re.search(r"epoch=(\d+)-", ckpt).group(1)
        epoch=int(epoch)
    else: 
        epoch=0
    mp['epoch']=epoch
    # return model.train(),mp
    return model.eval(),mp


def print_memory_usage():
    allocated_memory = torch.cuda.memory_allocated() / 1024**2  # Convert to MB
    reserved_memory = torch.cuda.memory_reserved() / 1024**2  # Convert to MB
    print(f"Allocated Memory: {allocated_memory:.2f} MB")
    print(f"Reserved Memory: {reserved_memory:.2f} MB")



def RRE(x, xinv):
    if isinstance(x, np.ndarray):
        return np.linalg.norm(x-xinv) / np.linalg.norm(x)
    elif isinstance(x, torch.Tensor):
        return torch.norm(x-xinv) / torch.norm(x)
    else:
        raise TypeError("Input should be a numpy array or a torch tensor")

def PSNR(xinv,x):
    if isinstance(x, np.ndarray):
        return 10 * np.log10(len(xinv) * np.max(xinv)**2 / np.linalg.norm(x-xinv)**2)
    elif isinstance(x, torch.Tensor):
        return 10 * torch.log10(xinv.numel() * torch.max(xinv)**2 / torch.norm(x-xinv)**2)
    else:
        raise TypeError("Input should be a numpy array or a torch tensor")


def r2_score(y_true, y_pred):
    y_true_mean = np.mean(y_true)
    ss_tot = np.sum((y_true - y_true_mean) ** 2)
    ss_res = np.sum((y_true - y_pred) ** 2)
    r2 = 1 - (ss_res / ss_tot)
    return r2


def count_psnr_rre(img,true=None,show=True,ssim_datarange=10,title=None):
    img=img.cpu().squeeze().numpy() if isinstance(img, torch.Tensor) else img
    true=true.cpu().squeeze().numpy() if isinstance(true, torch.Tensor) else true
    info={
        'PSNR':10 * np.log10(len(img.ravel()) * np.max(img.ravel())**2 / np.linalg.norm(true.ravel()-img.ravel())**2),
        'rre':np.linalg.norm(true-img) / np.linalg.norm(true),
        'SSIM':structural_similarity(true,img,data_range=true.max()),
        # 'nmse': (((true - img) ** 2)/(true.max()**2)).mean()*1000,
        'PCC':pearsonr(img.ravel(), true.ravel()).statistic,
        # 'r2':r2_score(true,img),
        'nmse':np.sum((true - img) ** 2) / np.sum(true ** 2),
        'mse':np.mean((true - img) ** 2)/img.size,
    }
    if show:
        if title is not None:
            print(title)
        for a in info:
            print(f"{a}:{info[a]:.4f}")
        print('')
    return info